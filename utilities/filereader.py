import base64
from pathlib import Path
from utilities.dependencies import logger 
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
from docx.table import Table

from docx.drawing import Drawing
from docx.image.image import Image
from docx.text.hyperlink import Hyperlink
import fitz  # PyMuPDF


class FileReader:    

    def __init__(self, file_pathes, include_images: bool = False):
        self.file_pathes = file_pathes
        self.include_images = include_images
        self.readers =  {
            '.txt': self._read_text,
            '.docx': self._read_docx,
            '.pdf': self._read_pdf,
            '.epub': self._read_epub,  # ebooklib
            '.rtf': self._read_rtf,    # striprtf
        }

    #dispatch method based on file extension
    # it chooses the appropriate method to read the file based on its extension
    # then reads all the data from it
    # some method should collect all the data and metadata to one single resulting dict
    def read(self):
        """
        Читає файли і повертає структуру:
        {
            "metadata": {
                "path": str,
                "type": str (txt/docx/pdf/etc),
                "size": int,
                "mtime": float
            },
            "content": {
                "pages": [...] для docx
                "text": str для txt
                ...
            }
        }
        """
        result = []
        for file_path in self.file_pathes:
            file_content = self.detector(file_path)
            size = Path(file_path).stat().st_size
            mtime = Path(file_path).stat().st_mtime
            
            result.append({
                "metadata": {
                    "path": str(file_path),
                    "type": file_path.suffix.lower().lstrip('.'),
                    "size": size,
                    "mtime": mtime
                },
                "content": file_content
            })
        return result
    
    # reads file extension and return reference to the function that can read it and call it
    def detector(self, path:Path):
        ext = path.suffix.lower()
        try:
            reader = self.readers.get(ext, self._read_text)  
        except Exception as e: # fallback на text
            logger.error(f"File with unsupported extension detected for {path}: {e}\n Falling back to text reader.")
            reader = self._read_text

        return reader(path)

    # adapter. Should turn results into a common format for all file types
    # what should this function be doing?
    def collect(self, file:dict):
        return file
    
    #after dispatching, call the appropriate method
    def _read_docx(self, file_path:Path):
        """Читає docx документ з витяганням тексту, картинок та гіперлінків.
        
        Повертає структуру:
        {
            "pages": [
                {
                    "number": 1,
                    "text": "текст з маркерами [[IMG:img_0]] [[LINK:link_0]]",
                    "media": [
                        {"kind": "image", "id": "img_0", "data": {...}},
                        {"kind": "link", "id": "link_0", "data": {...}}
                    ]
                }
            ]
        }
        """
        document = Document(file_path)
        pages = []
        pg_counter = 1
        current_page = []
        current_page_media = []
        img_counter = 0
        link_counter = 0

        def append_line(line: str) -> None:
            current_page.append(line)

        def append_media(obj: dict) -> None:
            current_page_media.append(obj)

        def flush_page() -> None:
            if current_page:
                pages.append({
                    "number": pg_counter,
                    "text": "\n".join(current_page),
                    "media": current_page_media[:]
                })
                current_page.clear()
                current_page_media.clear()

        def format_table(rows: list[list[str]]) -> list[str]:
            if not rows:
                return []
            col_count = max(len(row) for row in rows)
            normalized = [row + [""] * (col_count - len(row)) for row in rows]
            widths = [
                max(len(cell.replace("\n", " ").strip()) for cell in col)
                for col in zip(*normalized)
            ]

            def format_row(row: list[str]) -> str:
                padded = [cell.replace("\n", " ").strip().ljust(widths[i]) for i, cell in enumerate(row)]
                return "| " + " | ".join(padded) + " |"

            sep = "+" + "+".join("-" * (w + 2) for w in widths) + "+"
            output = [sep]
            for row in normalized:
                output.append(format_row(row))
                output.append(sep)
            return output
        
        for block in document.iter_inner_content():
            if isinstance(block, Paragraph):
                paragraph_text = ''
                has_page_break = False
                
                for run_or_hyperlink in block.iter_inner_content():
                    if isinstance(run_or_hyperlink, Run):
                        if run_or_hyperlink.contains_page_break:
                            has_page_break = True
                        for element in run_or_hyperlink.iter_inner_content():
                            if isinstance(element, Drawing):
                                if element.has_picture():
                                    image: Image = element.image
                                    img_id = f"img_{img_counter}"
                                    img_counter += 1
                                    
                                    paragraph_text += f"[[IMG:{img_id}]]"
                                    
                                    image_data = {
                                        "kind": "image",
                                        "id": img_id,
                                        "data": {
                                            "name": image.filename,
                                            "format": image.ext,
                                            "mime_type": image.content_type,
                                            "width_px": image.px_width,
                                            "height_px": image.px_height,
                                            "width_inches": float(image.width),
                                            "height_inches": float(image.height),
                                            "dpi_horizontal": image.horz_dpi,
                                            "dpi_vertical": image.vert_dpi,
                                            #"sha1": image.sha1
                                        }
                                    }

                                    if self.include_images:
                                        image_data["data"]["bytes_b64"] = base64.b64encode(image.blob).decode()

                                    append_media(image_data)
                                    
                            elif isinstance(element, str):
                                paragraph_text += element
                            else:
                                paragraph_text += str(element)
                                
                    elif isinstance(run_or_hyperlink, Hyperlink):
                        if run_or_hyperlink.contains_page_break:
                            has_page_break = True
                        link_id = f"link_{link_counter}"
                        link_counter += 1
                        
                        link_text = run_or_hyperlink.text
                        paragraph_text += f"{link_text}[[LINK:{link_id}]]"
                        
                        link_data = {
                            "kind": "link",
                            "id": link_id,
                            "data": {
                                "text": link_text,
                                "url": run_or_hyperlink.url,
                                "address": run_or_hyperlink.address,
                                "fragment": run_or_hyperlink.fragment,
                            }
                        }
                        append_media(link_data)
                    else:
                        paragraph_text += str(run_or_hyperlink)
                
                append_line(paragraph_text)
                if has_page_break:
                    flush_page()
                    pg_counter += 1

                
            elif isinstance(block, Table):
                rows = []
                for row in block.rows:
                    rows.append([cell.text for cell in row.cells])
                for line in format_table(rows):
                    append_line(line)
        
        flush_page()
        
        return {
            "pages": pages
        }

    def _read_pdf(self, file_path:Path):
        doc = fitz.open(file_path)
        pages = []
        img_counter = 0
        link_counter = 0
        table_counter = 0
        
        def format_table(rows: list[list[str]]) -> list[str]:
            """Format table in text representation, same as DOCX"""
            if not rows:
                return []
            col_count = max(len(row) for row in rows)
            normalized = [row + [""] * (col_count - len(row)) for row in rows]
            widths = [
                max(len(cell.replace("\n", " ").strip()) for cell in col)
                for col in zip(*normalized)
            ]

            def format_row(row: list[str]) -> str:
                padded = [cell.replace("\n", " ").strip().ljust(widths[i]) for i, cell in enumerate(row)]
                return "| " + " | ".join(padded) + " |"

            sep = "+" + "+".join("-" * (w + 2) for w in widths) + "+"
            output = [sep]
            for row in normalized:
                output.append(format_row(row))
                output.append(sep)
            return output
        
        for i in range(doc.page_count):
            page = doc.load_page(i)
            text_lines = []
            current_page_media = []
            
            # Get base text
            text = page.get_text()
            
            # Process tables
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    text_lines.append("=== Page contains tables ===\n")
                    for table_idx, table in enumerate(tables.tables):
                        table_id = f"table_{table_counter}"
                        table_counter += 1
                        
                        # Extract table data
                        table_data = table.extract()
                        if table_data:
                            # Add formatted table to text
                            text_lines.append(f"\n[[TABLE:{table_id}]]\n")
                            text_lines.extend(format_table(table_data))
                            text_lines.append("")
                            
                            # Store table metadata
                            current_page_media.append({
                                "kind": "table",
                                "id": table_id,
                                "data": {
                                    "bbox": list(table.bbox),
                                    "rows": len(table_data),
                                    "cols": len(table_data[0]) if table_data else 0,
                                    "cells": table_data
                                }
                            })
            except Exception as e:
                logger.error(f"Failed to extract tables from page {i}: {e}")
            
            #  main text
            text_lines.append(text)
            
            # image information
            image_info_list = page.get_image_info(hashes=True, xrefs=True)
            
            for img_info in image_info_list:
                img_id = f"img_{img_counter}"
                img_counter += 1
                
                image_data = {
                    "kind": "image",
                    "id": img_id,
                    "data": {
                        "xref": img_info.get("xref", 0),
                        "width": img_info.get("width", 0),
                        "height": img_info.get("height", 0),
                        "bpc": img_info.get("bpc", 8),
                        "colorspace": img_info.get("colorspace", 0),
                        "cs_name": img_info.get("cs-name", ""),
                        "xres": img_info.get("xres", 0),
                        "yres": img_info.get("yres", 0),
                        "size": img_info.get("size", 0),
                        "digest": img_info.get("digest", b"").hex() if img_info.get("digest") else "",
                        "bbox": list(img_info.get("bbox", (0, 0, 0, 0))),
                        "transform": list(img_info.get("transform", (1, 0, 0, 1, 0, 0))),
                        "has_mask": img_info.get("has-mask", False),
                    }
                }
                
                # actual image bytes
                if self.include_images and img_info.get("xref", 0) > 0:
                    try:
                        xref = img_info["xref"]
                        base_image = doc.extract_image(xref)
                        if base_image:
                            image_data["data"]["bytes_b64"] = base64.b64encode(base_image["image"]).decode()
                            image_data["data"]["format"] = base_image.get("ext", "")
                            image_data["data"]["mime_type"] = base_image.get("colorspace", "")
                    except Exception as e:
                        logger.error(f"Failed to extract image {xref} from page {i}: {e}")
                
                current_page_media.append(image_data)
                
                # image position
                bbox = img_info.get("bbox", (0, 0, 0, 0))
                marker = f"[[IMG:{img_id}]]"
                text_lines.append(f"\n{marker} at position {bbox}")
            
            links = page.get_links()
            for link in links:
                link_id = f"link_{link_counter}"
                link_counter += 1
                
                link_rect = link.get("from", [])
                link_uri = link.get("uri", "")
                link_type = link.get("kind", 0)
                
                link_data = {
                    "kind": "link",
                    "id": link_id,
                    "data": {
                        "uri": link_uri,
                        "type": link_type,
                        "rect": link_rect,
                        "page": link.get("page", -1),
                        "to": link.get("to", [])
                    }
                }
                current_page_media.append(link_data)
                
                # position info
                if link_uri:
                    marker = f"[[LINK:{link_id}]]"
                    text_lines.append(f"\n{marker} -> {link_uri} at {link_rect}")
            
            # all text
            final_text = "\n".join(text_lines)
            
            pages.append({
                "page": i + 1,
                "text": final_text,
                "media": current_page_media
            })
        
        metadata = doc.metadata
        doc.close()
        
        return { 
            "pages": pages,
            "pdf-metadata": metadata
        }

    def _read_text(self,file_path:Path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return text
    
    def _read_epub(self,file_path:Path):
        return {"text": "This is an epub file", "type": "epub"}
    
    def _read_rtf(self, file_path:Path):
        return {"text": "This is an rtf file", "type": "rtf"}
